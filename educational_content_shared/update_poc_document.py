#!/usr/bin/env python3
"""
Update Why_POC_Matters.docx with comprehensive content including:
- Visualization scenarios
- Real market examples
- Risk management rules
- Implementation guide
"""

import sys
from pathlib import Path

# Add project root to path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ python-docx not installed. Install with: pip install python-docx")

def create_poc_document_content():
    """Create comprehensive POC document content"""
    
    doc_content = {
        'title': 'Why POC Matters: Volume Profile Trading Guide',
        'sections': []
    }
    
    # Introduction
    doc_content['sections'].append({
        'heading': 'Introduction',
        'content': [
            'Point of Control (POC) is the price level where the most volume was traded during a session. It represents the fair value price where the most market participants agreed on value.',
            '',
            'Understanding POC helps you:',
            '• Identify key support/resistance levels',
            '• Find optimal entry and exit points',
            '• Understand market sentiment and fair value',
            '• Trade with institutional flow',
            ''
        ]
    })
    
    # Visualizations Section
    doc_content['sections'].append({
        'heading': 'Visualizations',
        'content': [
            'We have generated comprehensive volume profile visualizations using real market data from October 31, 2025. The Volume Profile charts provide superior understanding of price distribution and institutional trading activity.',
            '',
            'Available Chart Types:',
            '',
            '• Daily Volume Profile - Horizontal histogram showing volume distribution at each price level, clearly marking POC (Point of Control), Value Area High (VAH), and Value Area Low (VAL). This visualization gives the best understanding of where real volume occurred.',
            '',
            '• Multi-day POC Evolution - Tracks POC shifts across multiple trading sessions, showing how fair value changes over time.',
            '',
            'Generated visualizations include:',
            '',
            'Equities:',
            '• RELIANCE_2025-10-31_profile.png',
            '• HDFCBANK_2025-10-31_profile.png',
            '• TCS_2025-10-31_profile.png',
            '• INFY_2025-10-31_profile.png',
            '• ADANIPORTS_2025-10-31_profile.png',
            '• JSWSTEEL_2025-10-31_profile.png',
            '• M&M_2025-10-31_profile.png',
            '',
            'Futures:',
            '• BEL25NOVFUT_2025-10-31_profile.png',
            '• BANKNIFTY25NOV55200PE_2025-10-31_profile.png',
            '• NIFTY25NOV26050CE_2025-10-31_profile.png',
            '',
            'All charts are available in: educational_content_shared/poc_visualizations/',
            '',
            'Note: The Volume Profile charts (profile.png) are recommended as they provide clearer insight into volume distribution and POC compared to price-action charts.',
            ''
        ]
    })
    
    # Real Market Scenarios
    doc_content['sections'].append({
        'heading': 'Real Market Scenarios for Indian Stocks/Indices',
        'content': [
            'Scenario 1: NIFTY Range Day',
            '',
            'Setup: Price oscillates between Value Area boundaries',
            '',
            'Entry: Buy at LVA, Sell at UVA',
            '',
            'Stop: Beyond Value Area',
            '',
            'Target: Opposite boundary or POC',
            '',
            '---',
            '',
            'Scenario 2: BANKNIFTY Trend Day',
            '',
            'Setup: Price accepts above/below Value Area',
            '',
            'Entry: Pullback to POC with volume confirmation',
            '',
            'Stop: Beyond previous POC',
            '',
            'Target: Next volume node or extension',
            '',
            '---',
            '',
            'Scenario 3: Stock-Specific Breakout',
            '',
            'Setup: Consolidation with shrinking Value Area',
            '',
            'Entry: Break of LVN with increased volume',
            '',
            'Stop: Back into Value Area',
            '',
            'Target: Measured move or next POC level',
            ''
        ]
    })
    
    # Risk Management
    doc_content['sections'].append({
        'heading': 'Risk Management Rules',
        'content': [
            'Position Sizing:',
            '',
            'Calculate position size based on distance to nearest volume node:',
            '',
            '```python',
            '# Calculate position size based on distance to nearest volume node',
            'def calculate_position_size(capital, entry_price, stop_loss_price, risk_per_trade=1.0):',
            '    risk_per_share = abs(entry_price - stop_loss_price)',
            '    capital_risk = capital * (risk_per_trade / 100)',
            '    position_size = capital_risk / risk_per_share',
            '    return int(position_size)',
            '```',
            '',
            'Key Rules:',
            '',
            '• Never enter at POC without confirmation',
            '• Always place stops beyond significant volume nodes',
            '• Respect the Value Area - don\'t fight it',
            '• Volume precedes price - wait for confirmation',
            ''
        ]
    })
    
    # Common Mistakes
    doc_content['sections'].append({
        'heading': 'Common Mistakes to Avoid',
        'content': [
            '• Trading against the Volume Profile trend',
            '• Ignoring time-frame confluence',
            '• Placing stops at arbitrary levels',
            '• Overtrading during low-volume periods',
            '• Not adjusting for sector-specific volume patterns',
            ''
        ]
    })
    
    # Implementation
    doc_content['sections'].append({
        'heading': 'Action Steps for Implementation',
        'content': [
            '• Start with NIFTY 50 - most clean volume patterns',
            '• Use 30-minute charts for intraday and daily for swing',
            '• Paper trade for 2 weeks to identify reliable patterns',
            '• Focus on 2-3 high-volume stocks initially',
            '• Keep a trading journal specifically noting volume profile interactions',
            ''
        ]
    })
    
    # Advanced Tips
    doc_content['sections'].append({
        'heading': 'Advanced Tip: Session Analysis',
        'content': [
            'Break your analysis into:',
            '',
            'Opening Hour (9:15-10:30 AM) - Establishes initial value',
            '',
            'Mid Session (10:30 AM-2:30 PM) - Development of value',
            '',
            'Closing Hour (2:30-3:30 PM) - Final acceptance/rejection of value',
            ''
        ]
    })
    
    # Conclusion
    doc_content['sections'].append({
        'heading': 'Key Insight',
        'content': [
            'The key insight: Volume doesn\'t lie. While price can be manipulated in the short term, sustained volume at certain levels reveals true institutional interest. Master reading these footprints, and you\'ll be trading with the smart money.',
            ''
        ]
    })
    
    return doc_content


def update_docx_file():
    """Update the Word document with comprehensive content"""
    if not HAS_DOCX:
        print("❌ Cannot update Word document - python-docx not installed")
        print("   Install with: pip install python-docx")
        return False
    
    doc_path = project_root / "educational_content_shared" / "why_poc_matters.docx"
    
    # Create new document (overwrite existing)
    doc = Document()
    
    # Get content structure
    content = create_poc_document_content()
    
    # Add title
    title = doc.add_heading(content['title'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add sections
    for section in content['sections']:
        # Add heading
        heading = doc.add_heading(section['heading'], level=1)
        
        # Add content
        in_code_block = False
        code_lines = []
        
        for line in section['content']:
            if line.strip().startswith('```'):
                if in_code_block:
                    # End of code block - add formatted code
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run('\n'.join(code_lines))
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(10)
                    code_lines = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                continue
            
            if in_code_block:
                # Collect code lines
                code_lines.append(line)
                continue
            
            if line.startswith('•'):
                # Bullet point
                p = doc.add_paragraph(line, style='List Bullet')
            elif line.strip() == '':
                # Empty line
                doc.add_paragraph()
            elif line.strip().startswith('---'):
                # Horizontal rule
                p = doc.add_paragraph()
                p.add_run('─' * 50).bold = True
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)
    
    # Save document (overwrite original)
    doc.save(str(doc_path))
    print(f"✅ Updated document saved to: {doc_path}")
    print(f"   Document now includes all scenarios, visualizations, and implementation guide")
    return True


def create_text_version():
    """Create a plain text version that can be copied into Word"""
    content = create_poc_document_content()
    
    text_content = []
    text_content.append(content['title'])
    text_content.append('=' * len(content['title']))
    text_content.append('')
    
    for section in content['sections']:
        text_content.append(section['heading'])
        text_content.append('-' * len(section['heading']))
        text_content.append('')
        text_content.extend(section['content'])
        text_content.append('')
    
    text_path = project_root / "educational_content_shared" / "why_poc_matters_content.txt"
    with open(text_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(text_content))
    
    print(f"✅ Text content saved to: {text_path}")
    print("   You can copy this content into your Word document")
    return text_path


if __name__ == "__main__":
    print("Updating POC document...")
    
    # Try to update Word document if possible
    if HAS_DOCX:
        try:
            update_docx_file()
        except Exception as e:
            print(f"⚠️ Error updating Word document: {e}")
            print("   Creating text version instead...")
            create_text_version()
    else:
        # Create text version
        create_text_version()

